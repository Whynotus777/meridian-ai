"""Word document formatter for investment memos.

Takes a memo string + extracted_data dict and produces a professional
.docx file with:
  - Cover page with company name and key stats
  - Proper heading hierarchy (Heading 1 / Heading 2 styles)
  - Financial summary table (2-column metric/value)
  - Risk heat-map table (color-coded by severity)
  - Comparable companies table

Requires: python-docx
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Colour palette (RGB)
# ---------------------------------------------------------------------------
_NAVY       = RGBColor(0x1F, 0x4E, 0x79)
_DARK_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
_LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
_GREEN      = RGBColor(0x70, 0xAD, 0x47)
_AMBER      = RGBColor(0xFF, 0xC0, 0x00)
_RED        = RGBColor(0xFF, 0x00, 0x00)
_DARK_RED   = RGBColor(0xC0, 0x00, 0x00)
_ORANGE     = RGBColor(0xED, 0x7D, 0x31)
_BLACK      = RGBColor(0x00, 0x00, 0x00)

_SEVERITY_COLOURS: Dict[str, RGBColor] = {
    "Critical": _DARK_RED,
    "High":     _RED,
    "Medium":   _AMBER,
    "Low":      _GREEN,
}


# ---------------------------------------------------------------------------
# Low-level XML helpers (python-docx cell shading)
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, rgb: RGBColor):
    """Set a table cell's background fill via raw XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_colour = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _set_cell_borders(cell, colour_hex: str = "CCCCCC"):
    """Add thin borders to a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), colour_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _header_row(table, values: List[str],
                bg: RGBColor = None, text_colour: RGBColor = None):
    """Style the first row of a table as a header."""
    bg           = bg or _NAVY
    text_colour  = text_colour or _WHITE
    row          = table.rows[0]
    for cell, val in zip(row.cells, values):
        cell.text = val
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        run  = para.runs[0] if para.runs else para.add_run(val)
        run.text           = val
        run.font.bold      = True
        run.font.color.rgb = text_colour
        run.font.size      = Pt(10)
        para.alignment     = WD_ALIGN_PARAGRAPH.CENTER


def _data_row(table, row_idx: int, values: List[str],
              alt: bool = False, alignments: Optional[List] = None):
    """Write data into a row with alternating background."""
    bg = _LIGHT_GREY if alt else _WHITE
    row = table.rows[row_idx]
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = str(val) if val is not None else ""
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.runs[0].font.size = Pt(10) if para.runs else None
        if alignments and i < len(alignments):
            para.alignment = alignments[i]


# ---------------------------------------------------------------------------
# Document style helpers
# ---------------------------------------------------------------------------

def _set_heading_style(para, level: int = 1):
    """Apply Heading 1 or Heading 2 style."""
    style = f"Heading {level}"
    try:
        para.style = style
    except Exception:
        pass  # Style may not exist in minimal docx — fall back to manual
    if level == 1:
        for run in para.runs:
            run.font.color.rgb = _NAVY
            run.font.size      = Pt(14)
            run.font.bold      = True
    else:
        for run in para.runs:
            run.font.color.rgb = _DARK_BLUE
            run.font.size      = Pt(12)
            run.font.bold      = True


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal line (paragraph border)."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# ---------------------------------------------------------------------------
# Section parsers
# ---------------------------------------------------------------------------

_SECTION_PATTERNS = [
    r"^#+\s*(.+)",                       # markdown headings
    r"^(\d+\.\s+.+)$",                  # numbered: "1. EXECUTIVE SUMMARY"
    r"^([A-Z][A-Z\s&/:–-]{5,})$",       # ALL-CAPS lines
]


def _split_memo_sections(memo_text: str) -> List[Dict[str, str]]:
    """Split memo string into [{heading, body}] sections."""
    sections: List[Dict[str, str]] = []
    current_heading = "Investment Memo"
    current_body: List[str] = []

    for line in memo_text.splitlines():
        matched = False
        for pattern in _SECTION_PATTERNS:
            m = re.match(pattern, line.strip())
            if m:
                if current_body or current_heading:
                    sections.append({
                        "heading": current_heading,
                        "body": "\n".join(current_body).strip(),
                    })
                current_heading = m.group(1).strip().strip("#").strip()
                current_body = []
                matched = True
                break
        if not matched:
            current_body.append(line)

    if current_body or current_heading:
        sections.append({
            "heading": current_heading,
            "body": "\n".join(current_body).strip(),
        })

    return sections


# ---------------------------------------------------------------------------
# Table builders
# ---------------------------------------------------------------------------

def _add_financial_table(doc: Document, extracted_data: Dict[str, Any]):
    """Append a financial summary table to the document."""
    doc.add_heading("Financial Summary", level=2)

    fin = extracted_data.get("financials", {})
    rev = fin.get("revenue", {})
    ebitda = fin.get("ebitda", {})
    customers = extracted_data.get("customers", {})

    metrics = [
        ("LTM Revenue",                  _fmt_num(rev.get("ltm"))),
        ("Prior Year Revenue",           _fmt_num(rev.get("prior_year"))),
        ("Revenue CAGR (3yr)",           _fmt_pct(rev.get("cagr_3yr"))),
        ("LTM EBITDA",                   _fmt_num(ebitda.get("ltm"))),
        ("Adjusted EBITDA (LTM)",        _fmt_num(ebitda.get("adjusted_ebitda_ltm"))),
        ("EBITDA Margin",                _fmt_pct(ebitda.get("margin_ltm"))),
        ("Gross Margin",                 _fmt_pct(fin.get("gross_margin"))),
        ("Recurring Revenue %",          _fmt_pct(fin.get("recurring_revenue_pct"))),
        ("Total Customers",              _clean(customers.get("total_customers"))),
        ("Top Customer Concentration",   _fmt_pct(customers.get("top_customer_concentration"))),
        ("Annual Customer Retention",    _fmt_pct(customers.get("customer_retention"))),
        ("Net Revenue Retention",        _fmt_pct(customers.get("net_revenue_retention"))),
        ("Total Debt",                   _fmt_num(fin.get("debt"))),
        ("Cash & Equivalents",           _fmt_num(fin.get("cash"))),
    ]

    table = doc.add_table(rows=len(metrics) + 1, cols=2)
    table.style = "Table Grid"

    # Header row
    _header_row(table, ["Metric", "Value"])

    # Data rows
    for i, (metric, value) in enumerate(metrics):
        row = table.rows[i + 1]
        alt = (i % 2 == 1)

        # Metric label cell
        mc = row.cells[0]
        mc.text = metric
        _set_cell_bg(mc, _LIGHT_GREY if alt else _WHITE)
        _set_cell_borders(mc)
        if mc.paragraphs and mc.paragraphs[0].runs:
            mc.paragraphs[0].runs[0].font.bold = True
            mc.paragraphs[0].runs[0].font.size = Pt(10)

        # Value cell
        vc = row.cells[1]
        vc.text = value
        _set_cell_bg(vc, _LIGHT_GREY if alt else _WHITE)
        _set_cell_borders(vc)
        if vc.paragraphs and vc.paragraphs[0].runs:
            vc.paragraphs[0].runs[0].font.size = Pt(10)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Column widths
    for i, width in enumerate([Inches(3.0), Inches(3.0)]):
        for row in table.rows:
            row.cells[i].width = width

    doc.add_paragraph("")  # spacing


def _add_risk_table(doc: Document, risks: list):
    """Append a colour-coded risk register table."""
    doc.add_heading("Risk Register", level=2)

    if not risks:
        doc.add_paragraph("No risks identified.", style="Normal")
        return

    severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    sorted_risks = sorted(risks, key=lambda r: severity_order.get(r.severity, 4))

    cols = ["Category", "Severity", "Risk", "Description", "Mitigant"]
    table = doc.add_table(rows=len(sorted_risks) + 1, cols=len(cols))
    table.style = "Table Grid"
    _header_row(table, cols)

    for i, risk in enumerate(sorted_risks):
        row = table.rows[i + 1]
        values = [
            risk.category,
            risk.severity,
            risk.title,
            risk.description,
            risk.mitigant or "—",
        ]

        for j, (cell, val) in enumerate(zip(row.cells, values)):
            cell.text = str(val)
            _set_cell_borders(cell)

            if j == 1:  # Severity column — colour-coded
                colour = _SEVERITY_COLOURS.get(risk.severity, _LIGHT_GREY)
                _set_cell_bg(cell, colour)
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.bold      = True
                    p.runs[0].font.color.rgb = _WHITE
                    p.runs[0].font.size      = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                alt = (i % 2 == 1)
                _set_cell_bg(cell, _LIGHT_GREY if alt else _WHITE)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)

    # Set column widths
    col_widths = [Inches(1.0), Inches(0.9), Inches(1.5), Inches(2.5), Inches(1.5)]
    for row in table.rows:
        for j, width in enumerate(col_widths):
            if j < len(row.cells):
                row.cells[j].width = width

    doc.add_paragraph("")


def _add_comp_table(doc: Document, comps: list):
    """Append a comparable companies table."""
    doc.add_heading("Comparable Companies", level=2)

    if not comps:
        doc.add_paragraph("No comparable companies identified.", style="Normal")
        return

    cols = ["Company / Deal", "Type", "EV/Rev", "EV/EBITDA", "Rationale"]
    table = doc.add_table(rows=len(comps) + 1, cols=len(cols))
    table.style = "Table Grid"
    _header_row(table, cols)

    for i, comp in enumerate(comps):
        row = table.rows[i + 1]
        alt = (i % 2 == 1)
        values = [
            comp.name,
            comp.type.replace("_", " ").title(),
            f"{comp.ev_revenue:.1f}x" if comp.ev_revenue else "N/A",
            f"{comp.ev_ebitda:.1f}x" if comp.ev_ebitda else "N/A",
            comp.rationale[:120] + ("…" if len(comp.rationale) > 120 else ""),
        ]
        for j, (cell, val) in enumerate(zip(row.cells, values)):
            cell.text = str(val)
            _set_cell_bg(cell, _LIGHT_GREY if alt else _WHITE)
            _set_cell_borders(cell)
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            if j in (2, 3):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    col_widths = [Inches(1.6), Inches(1.4), Inches(0.8), Inches(1.0), Inches(2.6)]
    for row in table.rows:
        for j, width in enumerate(col_widths):
            if j < len(row.cells):
                row.cells[j].width = width

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def export_memo_docx(
    memo_text: str,
    extracted_data: Dict[str, Any],
    risks: list,
    comps: list,
    output_path: str,
) -> str:
    """Generate a formatted Word document from memo text and analysis data.

    Args:
        memo_text:      Raw memo string (markdown-ish) from MemoGenerator.
        extracted_data: Structured CIM dict from CIMExtractor.
        risks:          List of RiskFlag objects.
        comps:          List of Comparable objects.
        output_path:    Destination .docx path.

    Returns:
        Absolute path to the saved document.
    """
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = Document()
    _configure_document(doc)

    co = extracted_data.get("company_overview", {})
    company_name = co.get("company_name", "Unknown Company")
    industry     = f"{co.get('industry', 'N/A')} / {co.get('sub_industry', 'N/A')}"

    # ---- Cover page --------------------------------------------------------
    _add_cover(doc, company_name, industry, extracted_data)

    # ---- Memo sections -----------------------------------------------------
    sections = _split_memo_sections(memo_text)
    for section in sections:
        heading = section["heading"]
        body    = section["body"]

        # Determine heading level (numbered sections → H1, subsections → H2)
        level = 1 if re.match(r"^\d+\.", heading) or heading.isupper() else 2

        h = doc.add_heading(heading, level=level)
        _apply_heading_colour(h, level)

        if body:
            for para_text in body.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph(para_text)
                p.style = "Normal"
                for run in p.runs:
                    run.font.size = Pt(10)

    # ---- Structured tables -------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Appendix — Structured Analysis", level=1)

    _add_financial_table(doc, extracted_data)
    _add_horizontal_rule(doc)
    _add_risk_table(doc, risks or [])
    _add_horizontal_rule(doc)
    _add_comp_table(doc, comps or [])

    # ---- Footer metadata ---------------------------------------------------
    _add_footer(doc)

    doc.save(output_path)
    return os.path.abspath(output_path)


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------

def _configure_document(doc: Document):
    """Set page margins and default font."""
    from docx.oxml.ns import nsmap
    sections = doc.sections
    for section in sections:
        section.page_width   = Inches(8.5)
        section.page_height  = Inches(11)
        section.left_margin  = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin   = Inches(1.0)
        section.bottom_margin= Inches(1.0)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _apply_heading_colour(heading_para, level: int):
    """Apply brand colours to a heading paragraph's runs."""
    colour = _NAVY if level == 1 else _DARK_BLUE
    for run in heading_para.runs:
        run.font.color.rgb = colour
        run.font.size = Pt(14 if level == 1 else 12)


def _add_cover(
    doc: Document,
    company_name: str,
    industry: str,
    extracted_data: Dict[str, Any],
):
    """Build a styled cover page."""
    # Top banner paragraph
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run("CONFIDENTIAL INFORMATION MEMORANDUM")
    run.font.bold      = True
    run.font.size      = Pt(11)
    run.font.color.rgb = _DARK_BLUE
    run.font.all_caps  = True

    doc.add_paragraph("")

    # Company name
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(company_name)
    run.font.bold      = True
    run.font.size      = Pt(24)
    run.font.color.rgb = _NAVY

    # Industry line
    ind_para = doc.add_paragraph()
    ind_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ind_para.add_run(industry)
    run.font.size      = Pt(12)
    run.font.color.rgb = _DARK_BLUE

    doc.add_paragraph("")

    # Key stats mini-table
    fin = extracted_data.get("financials", {})
    rev = fin.get("revenue", {})
    ebitda = fin.get("ebitda", {})
    stats = [
        ("Revenue (LTM)",    _fmt_num(rev.get("ltm"))),
        ("EBITDA (LTM)",     _fmt_num(ebitda.get("ltm"))),
        ("EBITDA Margin",    _fmt_pct(ebitda.get("margin_ltm"))),
        ("Revenue CAGR",     _fmt_pct(rev.get("cagr_3yr"))),
    ]
    stats_table = doc.add_table(rows=1, cols=len(stats))
    stats_table.style = "Table Grid"
    row = stats_table.rows[0]
    for i, (label, val) in enumerate(stats):
        cell = row.cells[i]
        _set_cell_bg(cell, _NAVY)
        _set_cell_borders(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}\n")
        r1.font.bold      = True
        r1.font.color.rgb = _LIGHT_BLUE
        r1.font.size      = Pt(9)
        r2 = p.add_run(val)
        r2.font.bold      = True
        r2.font.color.rgb = _WHITE
        r2.font.size      = Pt(14)

    doc.add_paragraph("")

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(
        f"Analysis generated: {datetime.utcnow():%B %d, %Y}   |   "
        f"Powered by Meridian AI"
    )
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic    = True

    doc.add_page_break()


def _add_footer(doc: Document):
    """Add a confidentiality disclaimer paragraph."""
    doc.add_paragraph("")
    _add_horizontal_rule(doc)
    disclaimer = doc.add_paragraph(
        "CONFIDENTIAL — This document was generated by Meridian AI for internal "
        "analysis purposes only. Information is derived from publicly available CIM "
        "materials and should not be relied upon without independent verification. "
        f"Generated {datetime.utcnow():%Y-%m-%d %H:%M} UTC."
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in disclaimer.runs:
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.italic    = True


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _clean(val) -> str:
    if val is None or val == "not_provided":
        return "N/A"
    return str(val)


def _fmt_num(val) -> str:
    if val is None or val == "not_provided":
        return "N/A"
    try:
        v = float(val)
        if v >= 1_000_000:
            return f"${v / 1_000_000:.1f}M"
        if v >= 1_000:
            return f"${v / 1_000:.1f}K"
        return f"${v:,.0f}"
    except (ValueError, TypeError):
        return str(val)


def _fmt_pct(val) -> str:
    if val is None or val == "not_provided":
        return "N/A"
    try:
        return f"{float(val):.1%}"
    except (ValueError, TypeError):
        return str(val)
